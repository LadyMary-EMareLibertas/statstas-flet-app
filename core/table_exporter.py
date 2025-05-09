import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_apa_paragraph_style(paragraph):
    paragraph.paragraph_format.line_spacing = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

def set_cell_border_from_data(cell, border_top, border_bottom):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    def make_border_element(tag, border):
        el = OxmlElement(tag)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(border.get("thickness", 1) * 4))
        el.set(qn('w:color'), "FFFFFF" if border.get("color") == "white" else "000000")
        return el

    if border_top:
        tcBorders.append(make_border_element('w:top', border_top))
    if border_bottom:
        tcBorders.append(make_border_element('w:bottom', border_bottom))

    tcPr.append(tcBorders)

def export_table_to_word(table_data, filename="exports/table_export_test.docx"):
    print("🟢 Export started...")
    os.makedirs("exports", exist_ok=True)

    doc = Document()
    n_rows = len(table_data)
    n_cols = len(table_data[0]) if n_rows > 0 else 0

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = True

    merge_tracker = [[False]*n_cols for _ in range(n_rows)]

    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            if not cell.get("visible", True):
                merge_tracker[i][j] = True
                continue

            val = str(cell.get("value", ""))
            doc_cell = table.cell(i, j)
            doc_cell.text = val

            # 정렬
            if val.replace(".", "", 1).isdigit():
                doc_cell.paragraphs[0].alignment = 1  # CENTER
            else:
                doc_cell.paragraphs[0].alignment = 0  # LEFT

            doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_apa_paragraph_style(doc_cell.paragraphs[0])

            # 테두리 적용
            border_top = cell.get("border_top", {"color": "white", "thickness": 0})
            border_bottom = cell.get("border_bottom", {"color": "white", "thickness": 0})
            set_cell_border_from_data(doc_cell, border_top, border_bottom)

    # 병합 셀 처리
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            if cell.get("visible", True):
                span = 1
                for k in range(j + 1, n_cols):
                    if not table_data[i][k].get("visible", True):
                        span += 1
                    else:
                        break
                if span > 1:
                    merged_cell = table.cell(i, j)
                    for merge_j in range(1, span):
                        merged_cell = merged_cell.merge(table.cell(i, j + merge_j))
                        merge_tracker[i][j + merge_j] = True

    doc.save(filename)
    print(f"✅ Exported to {filename}")
    os.system(f"open '{filename}'")